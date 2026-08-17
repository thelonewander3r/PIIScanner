"""Optional Excel/Word/PDF adapter tests (piilint[office])."""

from __future__ import annotations

from pathlib import Path

import pytest
from typer.testing import CliRunner

pytest.importorskip("openpyxl")
pytest.importorskip("pypdf")
pytest.importorskip("docx")

from docx import Document
from openpyxl import load_workbook

from piilint.cli import app
from piilint.config import default_config
from piilint.engine import scan_path
from piilint.findings import EntityType
from piilint.redact import redact_tree

pytestmark = pytest.mark.office

CORPUS_XLSX = Path(__file__).resolve().parent.parent / "corpus" / "xlsx"
CORPUS_PDF = Path(__file__).resolve().parent.parent / "corpus" / "pdf"
CORPUS_DOCX = Path(__file__).resolve().parent.parent / "corpus" / "docx"

RAW = [
    "customer.alpha@retailmail.test",
    "ops.beta@corpmail.test",
    "234-56-7890",
    "512-48-3017",
    "4532015112830366",
]


def test_scan_xlsx_finds_planted() -> None:
    result = scan_path(CORPUS_XLSX / "customers.xlsx")
    entities = {f.entity for f in result.findings}
    assert EntityType.EMAIL in entities
    assert EntityType.SSN_US in entities
    blob = "\n".join(f.masked_sample for f in result.findings)
    for raw in RAW:
        assert raw not in blob


def test_scan_pdf_text_finds_planted() -> None:
    result = scan_path(CORPUS_PDF / "contacts.pdf")
    entities = {f.entity for f in result.findings}
    assert EntityType.EMAIL in entities
    assert EntityType.CREDIT_CARD in entities or EntityType.SSN_US in entities
    for f in result.findings:
        assert f.location.row is not None  # page number


def test_cli_scan_xlsx() -> None:
    runner = CliRunner()
    result = runner.invoke(app, ["scan", str(CORPUS_XLSX), "--fail-on", "high"])
    assert result.exit_code == 1, result.output
    for raw in RAW:
        assert raw not in result.output


def test_redact_xlsx_stretch(tmp_path: Path) -> None:
    out = tmp_path / "clean"
    result = redact_tree(CORPUS_XLSX / "customers.xlsx", out, config=default_config())
    assert result.files_written == 1
    dest = out / "customers.xlsx"
    assert dest.is_file()

    wb = load_workbook(dest, read_only=True, data_only=True)
    try:
        values = []
        for row in wb.active.iter_rows(values_only=True):
            values.extend(str(v) for v in row if v is not None)
        blob = "\n".join(values)
    finally:
        wb.close()
    for raw in RAW:
        assert raw not in blob

    wb2 = load_workbook(CORPUS_XLSX / "customers.xlsx", read_only=True, data_only=True)
    try:
        src_vals = []
        for row in wb2.active.iter_rows(values_only=True):
            src_vals.extend(str(v) for v in row if v is not None)
        assert "customer.alpha@retailmail.test" in "\n".join(src_vals)
    finally:
        wb2.close()


def test_scan_docx_finds_planted() -> None:
    result = scan_path(CORPUS_DOCX / "contacts.docx")
    entities = {f.entity for f in result.findings}
    assert EntityType.EMAIL in entities
    assert EntityType.SSN_US in entities
    blob = "\n".join(f.masked_sample for f in result.findings)
    for raw in RAW:
        assert raw not in blob


def test_cli_scan_docx() -> None:
    runner = CliRunner()
    result = runner.invoke(app, ["scan", str(CORPUS_DOCX), "--fail-on", "high"])
    assert result.exit_code == 1, result.output
    for raw in RAW:
        assert raw not in result.output


def test_redact_docx_stretch(tmp_path: Path) -> None:
    out = tmp_path / "clean"
    result = redact_tree(CORPUS_DOCX / "contacts.docx", out, config=default_config())
    assert result.files_written == 1
    dest = out / "contacts.docx"
    assert dest.is_file()

    doc = Document(str(dest))
    parts: list[str] = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    for section in doc.sections:
        parts.extend(p.text for p in section.header.paragraphs)
        parts.extend(p.text for p in section.footer.paragraphs)
    blob = "\n".join(parts)
    for raw in RAW:
        assert raw not in blob

    src = Document(str(CORPUS_DOCX / "contacts.docx"))
    src_blob = "\n".join(p.text for p in src.paragraphs)
    assert "customer.alpha@retailmail.test" in src_blob


def test_scan_docx_hard_negative_clean() -> None:
    result = scan_path(CORPUS_DOCX / "clean.docx")
    assert result.findings == []


NUMERIC_PHONES = ["2127350182", "4159032741", "4167350182"]


def test_redact_xlsx_numeric_phone_cells(tmp_path: Path) -> None:
    """Scan finds NANP phones stored as Excel numbers; redact masks them as text."""
    src = CORPUS_XLSX / "numeric_phones.xlsx"
    src_bytes = src.read_bytes()

    wb = load_workbook(src, data_only=True)
    try:
        ani = wb.active["A2"].value
        dnis = wb.active["B2"].value
        duration = wb.active["C2"].value
        assert isinstance(ani, (int, float)) and not isinstance(ani, bool)
        assert isinstance(dnis, (int, float)) and not isinstance(dnis, bool)
        assert isinstance(duration, (int, float)) and not isinstance(duration, bool)
        assert str(int(ani)) == "2127350182"
        assert str(int(dnis)) == "4159032741"
        assert int(duration) == 42
    finally:
        wb.close()

    before = scan_path(src)
    assert any(f.entity == EntityType.PHONE for f in before.findings)

    out = tmp_path / "clean"
    result = redact_tree(src, out, config=default_config())
    assert result.files_written == 1
    assert result.spans_redacted > 0
    dest = out / "numeric_phones.xlsx"
    assert dest.is_file()

    wb_out = load_workbook(dest, data_only=True)
    try:
        values: list[str] = []
        for row in wb_out.active.iter_rows():
            for cell in row:
                if cell.value is not None:
                    values.append(str(cell.value))
        blob = "\n".join(values)
        assert isinstance(wb_out.active["A2"].value, str)
        assert isinstance(wb_out.active["B2"].value, str)
        assert isinstance(wb_out.active["A3"].value, str)
        assert isinstance(wb_out.active["B3"].value, str)
        # Unchanged non-PII numbers stay numeric.
        assert isinstance(wb_out.active["C2"].value, (int, float))
        assert not isinstance(wb_out.active["C2"].value, bool)
        assert int(wb_out.active["C2"].value) == 42
    finally:
        wb_out.close()
    for raw in NUMERIC_PHONES:
        assert raw not in blob

    assert src.read_bytes() == src_bytes
    wb_src = load_workbook(src, data_only=True)
    try:
        assert wb_src.active["A2"].value == 2127350182
        assert wb_src.active["B2"].value == 4159032741
    finally:
        wb_src.close()

    after = scan_path(dest)
    after_blob = "\n".join(
        str(part) for f in after.findings for part in (f.masked_sample, f.normalized_value) if part
    )
    for raw in NUMERIC_PHONES:
        assert raw not in after_blob
    assert not any(f.entity == EntityType.PHONE for f in after.findings)
