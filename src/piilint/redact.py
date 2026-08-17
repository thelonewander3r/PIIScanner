"""Write cleaned file copies by rewriting PII spans (base wheel, no new deps).

Uses the same recognizers + ``mask_value`` as scan. Adapters/recognizers stay
untouched — orchestration lives here and in the CLI.
"""

from __future__ import annotations

import csv
import io
import json
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Any, cast

from piilint.adapters.text import TextAdapter, looks_binary
from piilint.config import Config
from piilint.engine import _build_registry, scan_text_matches
from piilint.findings import EntityType, Finding, Location, mask_value
from piilint.policy import apply_policy
from piilint.recognizers import Match, RecognizerRegistry
from piilint.walker import iter_files

# Supported formats for redact (Sprint 9 + 10 + office xlsx/docx/pdf).
_TEXT = TextAdapter()


class RedactError(ValueError):
    """Usage / safety error for the redact command (maps to exit 2)."""


@dataclass(frozen=True, slots=True)
class RedactResult:
    files_written: int
    files_skipped: int
    spans_redacted: int


@dataclass(frozen=True, slots=True)
class _Span:
    start: int
    end: int
    entity: EntityType
    value: str


def _supported(path: Path) -> str | None:
    """Return adapter kind or None if unsupported for redact."""
    suffix = path.suffix.lower()
    if suffix in {".csv", ".tsv"}:
        return "csv"
    if suffix in {".json", ".jsonl"}:
        return "json"
    if suffix == ".ipynb":
        return "notebook"
    if suffix == ".parquet":
        return "parquet"
    if suffix in {".xlsx", ".xlsm"}:
        return "xlsx"
    if suffix == ".docx":
        return "docx"
    if suffix == ".pdf":
        return "pdf"
    if _TEXT.supports(path):
        return "text"
    return None


def _line_index(text: str, offset: int) -> int:
    return text.count("\n", 0, offset) + 1


def _dedupe_spans(spans: list[_Span]) -> list[_Span]:
    """Keep non-overlapping spans (earlier start wins; ties → longer span)."""
    ordered = sorted(spans, key=lambda s: (s.start, -(s.end - s.start)))
    kept: list[_Span] = []
    cursor = -1
    for span in ordered:
        if span.start < cursor:
            continue
        if span.end <= span.start:
            continue
        kept.append(span)
        cursor = span.end
    return kept


def apply_span_replacements(text: str, spans: list[_Span]) -> str:
    """Replace spans with ``mask_value`` results (reverse offset order)."""
    kept = _dedupe_spans(spans)
    out = text
    for span in sorted(kept, key=lambda s: s.start, reverse=True):
        replacement = mask_value(span.value, span.entity)
        out = out[: span.start] + replacement + out[span.end :]
    return out


def _filter_matches(
    text: str,
    matches: list[Match],
    *,
    config: Config,
    rel_path: str,
    context_key: str | None = None,
) -> list[_Span]:
    """Convert matches ? findings, apply policy, return surviving spans."""
    if not matches:
        return []
    findings: list[Finding] = []
    line_texts: dict[tuple[str, int], str] = {}
    for idx, line in enumerate(text.split("\n"), start=1):
        line_texts[(rel_path, idx)] = line

    for i, match in enumerate(matches):
        line_no = _line_index(text, match.start)
        findings.append(
            Finding.create(
                entity=match.entity,
                raw_value=match.value,
                location=Location(
                    path=rel_path,
                    line=line_no,
                    column=context_key,
                    offset=match.start,
                ),
                confidence=match.confidence,
                severity=match.severity,
                occurrence_index=i,
            )
        )
    kept = apply_policy(findings, config, line_texts=line_texts)
    offset_entity = {(f.location.offset, f.entity) for f in kept if f.location.offset is not None}
    return [
        _Span(start=m.start, end=m.end, entity=m.entity, value=m.value)
        for m in matches
        if (m.start, m.entity) in offset_entity
    ]


def redact_plain_text(
    text: str,
    *,
    registry: RecognizerRegistry,
    config: Config,
    rel_path: str,
    context_key: str | None = None,
) -> tuple[str, int]:
    """Redact PII spans in a plain string. Returns (new_text, span_count)."""
    matches = scan_text_matches(
        text,
        registry.enabled_recognizers(),
        context_key=context_key,
        min_confidence=config.scan.min_confidence,
    )
    spans = _filter_matches(
        text, matches, config=config, rel_path=rel_path, context_key=context_key
    )
    if not spans:
        return text, 0
    return apply_span_replacements(text, spans), len(_dedupe_spans(spans))


def _read_text_file(path: Path) -> tuple[str, bool]:
    """Return (text, had_bom). Normalizes newlines to \\n for scanning; caller may rewrite."""
    data = path.read_bytes()
    had_bom = data.startswith(b"\xef\xbb\xbf")
    if had_bom:
        data = data[3:]
    if looks_binary(data[:8192]):
        raise RedactError(f"Refusing binary file: {path}")
    text = data.decode("utf-8", errors="replace")
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    return text, had_bom


def _write_text(path: Path, text: str, *, had_bom: bool = False) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    payload = text
    data = payload.encode("utf-8")
    if had_bom:
        data = b"\xef\xbb\xbf" + data
    path.write_bytes(data)


def _redact_json_value(
    value: Any,
    *,
    registry: RecognizerRegistry,
    config: Config,
    rel_path: str,
    key: str | None,
) -> tuple[Any, int]:
    count = 0
    if isinstance(value, dict):
        out: dict[Any, Any] = {}
        for child_key, child in value.items():
            new_child, n = _redact_json_value(
                child,
                registry=registry,
                config=config,
                rel_path=rel_path,
                key=str(child_key),
            )
            out[child_key] = new_child
            count += n
        return out, count
    if isinstance(value, list):
        items = []
        for item in value:
            new_item, n = _redact_json_value(
                item, registry=registry, config=config, rel_path=rel_path, key=key
            )
            items.append(new_item)
            count += n
        return items, count
    if value is None or isinstance(value, (int, float, bool)):
        return value, 0
    text = str(value)
    if not text.strip():
        return value, 0
    redacted, n = redact_plain_text(
        text,
        registry=registry,
        config=config,
        rel_path=rel_path,
        context_key=key,
    )
    return redacted, n


def _redact_json_file(
    path: Path,
    *,
    registry: RecognizerRegistry,
    config: Config,
    rel_path: str,
) -> tuple[str, int]:
    raw = path.read_text(encoding="utf-8", errors="replace")
    if path.suffix.lower() == ".jsonl":
        lines_out: list[str] = []
        total = 0
        for line in raw.splitlines():
            if not line.strip():
                lines_out.append(line)
                continue
            try:
                data = json.loads(line)
            except json.JSONDecodeError:
                redacted, n = redact_plain_text(
                    line, registry=registry, config=config, rel_path=rel_path
                )
                lines_out.append(redacted)
                total += n
                continue
            new_data, n = _redact_json_value(
                data, registry=registry, config=config, rel_path=rel_path, key=None
            )
            lines_out.append(json.dumps(new_data, ensure_ascii=False, separators=(",", ":")))
            total += n
        # Preserve trailing newline if original had one
        body = "\n".join(lines_out)
        if raw.endswith("\n"):
            body += "\n"
        return body, total

    data = json.loads(raw)
    new_data, n = _redact_json_value(
        data, registry=registry, config=config, rel_path=rel_path, key=None
    )
    return json.dumps(new_data, ensure_ascii=False, indent=2) + "\n", n


def _redact_csv_file(
    path: Path,
    *,
    registry: RecognizerRegistry,
    config: Config,
    rel_path: str,
) -> tuple[str, int]:
    delimiter = "\t" if path.suffix.lower() == ".tsv" else ","
    text = path.read_text(encoding="utf-8", errors="replace")
    # Strip BOM for parsing
    if text.startswith("\ufeff"):
        text = text[1:]
    reader = csv.reader(io.StringIO(text), delimiter=delimiter)
    rows = list(reader)
    if not rows:
        return text, 0
    headers = rows[0]
    total = 0
    out_rows: list[list[str]] = [headers]
    for row in rows[1:]:
        new_row: list[str] = []
        for idx, cell in enumerate(row):
            header = headers[idx] if idx < len(headers) else None
            if not cell:
                new_row.append(cell)
                continue
            redacted, n = redact_plain_text(
                cell,
                registry=registry,
                config=config,
                rel_path=rel_path,
                context_key=header,
            )
            new_row.append(redacted)
            total += n
        out_rows.append(new_row)
    buf = io.StringIO()
    writer = csv.writer(buf, delimiter=delimiter, lineterminator="\n")
    writer.writerows(out_rows)
    return buf.getvalue(), total


def _nb_text(value: Any) -> str:
    if isinstance(value, list):
        return "".join(str(part) for part in value)
    if value is None:
        return ""
    return value if isinstance(value, str) else str(value)


def _redact_notebook_file(
    path: Path,
    dest: Path,
    *,
    registry: RecognizerRegistry,
    config: Config,
    rel_path: str,
) -> int:
    """Redact source + scanned output text in a notebook; write nbformat-compatible copy."""
    import nbformat
    from nbformat.reader import NotJSONError

    try:
        nb = nbformat.read(path, as_version=4)  # type: ignore[no-untyped-call]
    except (OSError, NotJSONError, nbformat.ValidationError, ValueError) as exc:
        raise RedactError(f"Invalid notebook {rel_path}: {exc}") from exc

    total = 0
    for cell in nb.cells:
        source = _nb_text(cell.get("source"))
        new_source, n = redact_plain_text(
            source, registry=registry, config=config, rel_path=rel_path
        )
        cell["source"] = new_source
        total += n

        if cell.get("cell_type") != "code":
            continue
        outputs = cell.get("outputs") or []
        for output in outputs:
            output_type = getattr(output, "output_type", None) or (
                output.get("output_type") if isinstance(output, dict) else None
            )
            if output_type == "stream":
                raw = getattr(output, "text", None)
                if raw is None and isinstance(output, dict):
                    raw = output.get("text")
                block = _nb_text(raw)
                if not block:
                    continue
                redacted, n = redact_plain_text(
                    block, registry=registry, config=config, rel_path=rel_path
                )
                if isinstance(output, dict):
                    output["text"] = redacted
                else:
                    output.text = redacted
                total += n
            elif output_type in {"execute_result", "display_data"}:
                data = getattr(output, "data", None)
                if data is None and isinstance(output, dict):
                    data = output.get("data")
                if not isinstance(data, dict) or "text/plain" not in data:
                    continue
                plain = data["text/plain"]
                block = _nb_text(plain)
                if not block:
                    continue
                redacted, n = redact_plain_text(
                    block, registry=registry, config=config, rel_path=rel_path
                )
                data["text/plain"] = redacted
                total += n
            # Binary / image payloads left untouched (scan does not extract them as text).

    dest.parent.mkdir(parents=True, exist_ok=True)
    nbformat.write(nb, dest)  # type: ignore[no-untyped-call]
    return total


def _is_stringish_arrow_type(arrow_type: Any) -> bool:
    import pyarrow as pa

    if pa.types.is_string(arrow_type) or pa.types.is_large_string(arrow_type):
        return True
    if pa.types.is_dictionary(arrow_type):
        value_type = arrow_type.value_type
        return bool(pa.types.is_string(value_type) or pa.types.is_large_string(value_type))
    return False


def _redact_parquet_file(
    path: Path,
    dest: Path,
    *,
    registry: RecognizerRegistry,
    config: Config,
    rel_path: str,
) -> int:
    """Redact string / dictionary-string columns; leave other types unchanged.

    Nested/list/struct columns are not rewritten in v1 (documented limitation).
    """
    import pyarrow as pa
    import pyarrow.parquet as pq

    try:
        table = pq.read_table(path)
    except (OSError, pa.ArrowInvalid, pa.ArrowTypeError) as exc:
        raise RedactError(f"Invalid parquet {rel_path}: {exc}") from exc

    total = 0
    arrays: list[Any] = []
    names: list[str] = []
    for i, field in enumerate(table.schema):
        names.append(field.name)
        col = table.column(i)
        if not _is_stringish_arrow_type(field.type):
            arrays.append(col)
            continue
        values = col.to_pylist()
        new_values: list[str | None] = []
        for value in values:
            if value is None:
                new_values.append(None)
                continue
            text = value if isinstance(value, str) else str(value)
            if not text:
                new_values.append(text)
                continue
            redacted, n = redact_plain_text(
                text,
                registry=registry,
                config=config,
                rel_path=rel_path,
                context_key=field.name,
            )
            new_values.append(redacted)
            total += n
        # Normalize to plain string column (dictionary ? string after rewrite).
        arrays.append(pa.array(new_values, type=pa.string()))

    new_table = pa.table({name: arrays[idx] for idx, name in enumerate(names)})
    dest.parent.mkdir(parents=True, exist_ok=True)
    try:
        pq.write_table(new_table, dest)
    except (OSError, pa.ArrowInvalid, pa.ArrowTypeError) as exc:
        raise RedactError(f"Failed to write parquet {rel_path}: {exc}") from exc
    return total


def _redact_xlsx_file(
    path: Path,
    dest: Path,
    *,
    registry: RecognizerRegistry,
    config: Config,
    rel_path: str,
) -> int:
    """Rewrite string and numeric PII cells into a new workbook under ``dest``."""
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise RedactError(
            'Excel redact requires piilint[office]. Install with: pip install "piilint[office]"'
        ) from exc

    try:
        wb = load_workbook(path)
    except OSError as exc:
        raise RedactError(f"Invalid workbook {rel_path}: {exc}") from exc

    total = 0
    for sheet in wb.worksheets:
        headers: dict[int, str] = {}
        for row in sheet.iter_rows():
            for cell in row:
                value = cell.value
                if value is None or isinstance(value, bool):
                    # Keep empty/bools unchanged. Numeric PII is stringified below.
                    continue
                if isinstance(value, float) and value != value:  # NaN
                    continue
                text = str(value).strip()
                if not text:
                    continue
                if cell.row == 1:
                    headers[cell.column] = text
                context = headers.get(cell.column) or sheet.title
                redacted, n = redact_plain_text(
                    text,
                    registry=registry,
                    config=config,
                    rel_path=rel_path,
                    context_key=context,
                )
                if n:
                    # Mask as text — a numeric phone cannot stay a number.
                    cell.value = redacted
                    total += n

    dest.parent.mkdir(parents=True, exist_ok=True)
    try:
        wb.save(dest)
    except OSError as exc:
        raise RedactError(f"Failed to write workbook {rel_path}: {exc}") from exc
    return total


def _redact_paragraphs(
    paragraphs: object,
    *,
    registry: RecognizerRegistry,
    config: Config,
    rel_path: str,
    context_key: str,
) -> int:
    total = 0
    for paragraph in paragraphs:  # type: ignore[attr-defined]
        text = paragraph.text or ""
        if not text.strip():
            continue
        redacted, n = redact_plain_text(
            text,
            registry=registry,
            config=config,
            rel_path=rel_path,
            context_key=context_key,
        )
        if n:
            paragraph.text = redacted
            total += n
    return total


def _redact_docx_tables(
    tables: object,
    *,
    registry: RecognizerRegistry,
    config: Config,
    rel_path: str,
    context_prefix: str,
) -> int:
    """Redact via cell.text; first-row headers become context keys (like xlsx)."""
    total = 0
    table_list = list(cast(Any, tables))
    for table_idx, table in enumerate(table_list, start=1):
        label = f"{context_prefix}:table{table_idx}" if context_prefix else f"table{table_idx}"
        headers: dict[int, str] = {}
        for row_idx, row in enumerate(table.rows):
            # Dedupe only within a row: horizontally merged cells repeat the same tc.
            # Do NOT keep a table-wide seen set — python-docx may recycle cell wrappers.
            seen_row: set[int] = set()
            for col_idx, cell in enumerate(row.cells):
                cell_id = id(cell._tc)  # noqa: SLF001 - oxml tc element
                if cell_id in seen_row:
                    continue
                seen_row.add(cell_id)
                raw = cell.text or ""
                if not raw.strip():
                    continue
                if row_idx == 0:
                    headers[col_idx] = raw.strip()
                context = headers.get(col_idx) or label
                redacted, n = redact_plain_text(
                    raw,
                    registry=registry,
                    config=config,
                    rel_path=rel_path,
                    context_key=context,
                )
                if n:
                    cell.text = redacted
                    total += n
    return total


def _redact_docx_file(
    path: Path,
    dest: Path,
    *,
    registry: RecognizerRegistry,
    config: Config,
    rel_path: str,
) -> int:
    """Rewrite paragraph/table/header/footer text into a new .docx under ``dest``."""
    try:
        from docx import Document
        from docx.opc.exceptions import PackageNotFoundError
    except ImportError as exc:
        raise RedactError(
            'Word redact requires piilint[office]. Install with: pip install "piilint[office]"'
        ) from exc

    try:
        document = Document(str(path))
    except (OSError, PackageNotFoundError, ValueError, KeyError) as exc:
        raise RedactError(f"Invalid document {rel_path}: {exc}") from exc

    total = 0
    total += _redact_paragraphs(
        document.paragraphs,
        registry=registry,
        config=config,
        rel_path=rel_path,
        context_key="body",
    )
    total += _redact_docx_tables(
        document.tables,
        registry=registry,
        config=config,
        rel_path=rel_path,
        context_prefix="",
    )
    for section_idx, section in enumerate(document.sections, start=1):
        for part_name, container in (
            ("header", section.header),
            ("footer", section.footer),
        ):
            context = f"section{section_idx}:{part_name}"
            total += _redact_paragraphs(
                container.paragraphs,
                registry=registry,
                config=config,
                rel_path=rel_path,
                context_key=context,
            )
            total += _redact_docx_tables(
                container.tables,
                registry=registry,
                config=config,
                rel_path=rel_path,
                context_prefix=context,
            )

    dest.parent.mkdir(parents=True, exist_ok=True)
    try:
        document.save(str(dest))
    except OSError as exc:
        raise RedactError(f"Failed to write document {rel_path}: {exc}") from exc
    return total


_PDF_TEXT_OPS = {b"Tj", b"TJ", b"'", b'"'}


def _pdf_replace_in_operand(value: object, replacements: list[tuple[str, str]]) -> object:
    """Replace raw PII substrings in a PDF text operand; leave numbers untouched."""
    if isinstance(value, (int, float)):
        return value
    if isinstance(value, bytes):
        try:
            text = value.decode("latin-1")
        except UnicodeDecodeError:
            return value
        new = text
        for raw, masked in replacements:
            if raw and raw in new:
                new = new.replace(raw, masked)
        if new == text:
            return value
        from pypdf.generic import create_string_object

        return create_string_object(new)
    if isinstance(value, str):
        new = value
        for raw, masked in replacements:
            if raw and raw in new:
                new = new.replace(raw, masked)
        if new == value:
            return value
        from pypdf.generic import create_string_object

        return create_string_object(new)
    return value


def _pdf_content_replace(page: Any, reader: Any, replacements: list[tuple[str, str]]) -> bool:
    """Replace raw match strings in text-showing operators. Returns True if any changed."""
    from pypdf.generic import ArrayObject, ContentStream, NameObject

    contents = page.get_contents()
    if contents is None:
        return False
    stream = ContentStream(contents, reader)
    changed = False
    new_ops: list[tuple[list[Any], bytes]] = []
    for operands, operator in stream.operations:
        ops = list(operands)
        if operator in _PDF_TEXT_OPS and ops:
            if operator == b"TJ":
                items = list(ops[0])
                new_items = [_pdf_replace_in_operand(item, replacements) for item in items]
                if new_items != items:
                    ops[0] = ArrayObject(new_items)
                    changed = True
            else:
                idx = len(ops) - 1
                new_val = _pdf_replace_in_operand(ops[idx], replacements)
                if new_val != ops[idx]:
                    ops[idx] = new_val
                    changed = True
        new_ops.append((ops, operator))
    if changed:
        stream.operations = new_ops
        page[NameObject("/Contents")] = stream
    return changed


def _pdf_extract_text(page: Any) -> str:
    try:
        text = page.extract_text() or ""
    except Exception:  # noqa: BLE001 — malformed page content
        return ""
    return text.replace("\r\n", "\n").replace("\r", "\n")


def _pdf_page_still_has_raw(page: Any, replacements: list[tuple[str, str]]) -> bool:
    text = _pdf_extract_text(page)
    return any(raw in text for raw, _masked in replacements if raw)


def _pdf_ensure_font(page: Any) -> None:
    from pypdf.generic import DictionaryObject, NameObject

    resources = page.get("/Resources")
    if resources is None:
        resources = DictionaryObject()
        page[NameObject("/Resources")] = resources
    else:
        resources = resources.get_object()
    fonts = resources.get("/Font")
    if fonts is None:
        fonts = DictionaryObject()
        resources[NameObject("/Font")] = fonts
    else:
        fonts = fonts.get_object()
    if "/F1" in fonts or NameObject("/F1") in fonts:
        return
    fonts[NameObject("/F1")] = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )


def _pdf_reconstruct_page(page: Any, reader: Any, text: str) -> None:
    """Rebuild a page content stream from redacted extracted text (layout may change)."""
    from pypdf.generic import ContentStream, NameObject, NumberObject, create_string_object

    lines = text.split("\n")
    operations: list[tuple[list[Any], bytes]] = [
        ([], b"BT"),
        ([NameObject("/F1"), NumberObject(12)], b"Tf"),
        ([NumberObject(50), NumberObject(750)], b"Td"),
    ]
    for index, line in enumerate(lines):
        if index:
            operations.append(([NumberObject(0), NumberObject(-16)], b"Td"))
        operations.append(([create_string_object(line)], b"Tj"))
    operations.append(([], b"ET"))
    contents = page.get_contents()
    stream = (
        ContentStream(contents, reader) if contents is not None else ContentStream(None, reader)
    )
    stream.operations = operations
    page[NameObject("/Contents")] = stream
    _pdf_ensure_font(page)


def _pdf_spans_for_text(
    text: str,
    *,
    registry: RecognizerRegistry,
    config: Config,
    rel_path: str,
    context_key: str,
) -> list[_Span]:
    matches = scan_text_matches(
        text,
        registry.enabled_recognizers(),
        context_key=context_key,
        min_confidence=config.scan.min_confidence,
    )
    return _dedupe_spans(
        _filter_matches(text, matches, config=config, rel_path=rel_path, context_key=context_key)
    )


def _redact_pdf_file(
    path: Path,
    dest: Path,
    *,
    registry: RecognizerRegistry,
    config: Config,
    rel_path: str,
) -> int:
    """Rewrite embedded-text PII into a new PDF under ``dest`` (pypdf only; no OCR)."""
    try:
        from pypdf import PdfReader, PdfWriter
        from pypdf.errors import PdfReadError
    except ImportError as exc:
        raise RedactError(
            'PDF redact requires piilint[office]. Install with: pip install "piilint[office]"'
        ) from exc

    try:
        reader = PdfReader(str(path))
    except (OSError, PdfReadError, ValueError) as exc:
        raise RedactError(f"Invalid PDF {rel_path}: {exc}") from exc

    writer = PdfWriter()
    total = 0
    any_text = False
    for page_idx, page in enumerate(reader.pages, start=1):
        text = _pdf_extract_text(page)
        if text.strip():
            any_text = True
        spans = _pdf_spans_for_text(
            text,
            registry=registry,
            config=config,
            rel_path=rel_path,
            context_key=f"page{page_idx}",
        )
        if not spans:
            writer.add_page(page)
            continue
        replacements = sorted(
            {(span.value, mask_value(span.value, span.entity)) for span in spans},
            key=lambda pair: len(pair[0]),
            reverse=True,
        )
        replaced = _pdf_content_replace(page, reader, replacements)
        if (not replaced) or _pdf_page_still_has_raw(page, replacements):
            redacted = apply_span_replacements(text, spans)
            _pdf_reconstruct_page(page, reader, redacted)
        total += len(spans)
        writer.add_page(page)

    if not any_text:
        print(
            f"piilint: no embedded text in {rel_path} — PDF redact is a no-op (no OCR)",
            file=sys.stderr,
        )

    dest.parent.mkdir(parents=True, exist_ok=True)
    try:
        with dest.open("wb") as handle:
            writer.write(handle)
    except OSError as exc:
        raise RedactError(f"Failed to write PDF {rel_path}: {exc}") from exc
    return total


def _safe_out_path(out_root: Path, rel: str) -> Path:
    """Resolve destination under out_root; refuse escapes."""
    # Normalize rel to forbid absolute / drive / .. escape
    rel_path = Path(rel)
    if rel_path.is_absolute() or bool(rel_path.drive):
        raise RedactError(f"Refusing absolute relative path: {rel}")
    dest = (out_root / rel_path).resolve()
    try:
        dest.relative_to(out_root.resolve())
    except ValueError as exc:
        raise RedactError(f"Refusing path outside output dir: {rel}") from exc
    return dest


def redact_tree(
    target: Path,
    output_dir: Path,
    *,
    config: Config,
    enable_ner: bool = False,
) -> RedactResult:
    """Redact supported files under ``target`` into ``output_dir`` (mirrors relpaths)."""
    if not target.exists():
        raise RedactError(f"Path not found: {target}")
    out_root = output_dir.resolve()
    out_root.mkdir(parents=True, exist_ok=True)

    root = target.resolve()
    base = root if root.is_dir() else root.parent
    registry = _build_registry(config, enable_ner=enable_ner)

    written = 0
    skipped = 0
    spans = 0

    for file_path in iter_files(root, exclude=config.scan.exclude or None):
        kind = _supported(file_path)
        if kind is None:
            skipped += 1
            continue
        try:
            rel = file_path.relative_to(base).as_posix()
        except ValueError:
            rel = file_path.name
        dest = _safe_out_path(out_root, rel)

        try:
            if kind == "text":
                text, had_bom = _read_text_file(file_path)
                new_text, n = redact_plain_text(
                    text, registry=registry, config=config, rel_path=rel
                )
                _write_text(dest, new_text, had_bom=had_bom)
            elif kind == "json":
                new_text, n = _redact_json_file(
                    file_path, registry=registry, config=config, rel_path=rel
                )
                _write_text(dest, new_text)
            elif kind == "csv":
                new_text, n = _redact_csv_file(
                    file_path, registry=registry, config=config, rel_path=rel
                )
                _write_text(dest, new_text)
            elif kind == "notebook":
                n = _redact_notebook_file(
                    file_path, dest, registry=registry, config=config, rel_path=rel
                )
            elif kind == "parquet":
                n = _redact_parquet_file(
                    file_path, dest, registry=registry, config=config, rel_path=rel
                )
            elif kind == "xlsx":
                n = _redact_xlsx_file(
                    file_path, dest, registry=registry, config=config, rel_path=rel
                )
            elif kind == "docx":
                n = _redact_docx_file(
                    file_path, dest, registry=registry, config=config, rel_path=rel
                )
            elif kind == "pdf":
                n = _redact_pdf_file(
                    file_path, dest, registry=registry, config=config, rel_path=rel
                )
            else:
                skipped += 1
                continue
        except (OSError, json.JSONDecodeError, csv.Error, UnicodeError) as exc:
            raise RedactError(f"Failed to redact {rel}: {exc}") from exc

        written += 1
        spans += n

    return RedactResult(files_written=written, files_skipped=skipped, spans_redacted=spans)
