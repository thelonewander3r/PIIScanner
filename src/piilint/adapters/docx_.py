"""Word (.docx) adapter — optional ``piilint[office]`` (python-docx). No legacy .doc."""

from __future__ import annotations

import importlib.util
import sys
from collections.abc import Iterator
from pathlib import Path

from piilint.adapters import Unit

_INSTALL_HINT = 'pip install "piilint[office]"'
_warned = False


def office_docx_available() -> bool:
    return importlib.util.find_spec("docx") is not None


def _warn_missing_once() -> None:
    global _warned
    if _warned:
        return
    _warned = True
    print(
        f"piilint: skipping .docx (python-docx not installed). Install with: {_INSTALL_HINT}",
        file=sys.stderr,
    )


def _iter_paragraph_units(
    paragraphs: object,
    *,
    rel_path: str,
    context_key: str,
    line_start: int,
) -> Iterator[tuple[int, Unit]]:
    """Yield (next_line, unit) for nonempty paragraph texts."""
    line = line_start
    for para in paragraphs:  # type: ignore[attr-defined]
        text = (para.text or "").replace("\r\n", "\n").replace("\r", "\n")
        if not text.strip():
            continue
        yield (
            line,
            Unit(
                text=text,
                path=rel_path,
                line=line,
                context_key=context_key,
                aggregate=False,
            ),
        )
        line += 1


def _iter_table_units(
    table: object,
    *,
    rel_path: str,
    table_label: str,
    line_start: int,
) -> Iterator[tuple[int, Unit]]:
    """Yield units for table cells; first row supplies column context keys."""
    line = line_start
    headers: dict[int, str] = {}
    seen: set[object] = set()
    for row_idx, row in enumerate(table.rows):  # type: ignore[attr-defined]
        for col_idx, cell in enumerate(row.cells):
            tc = cell._tc  # noqa: SLF001 — unique cell element (keep ref)
            if tc in seen:
                continue
            seen.add(tc)
            text = (cell.text or "").replace("\r\n", "\n").replace("\r", "\n")
            if not text.strip():
                continue
            if row_idx == 0:
                headers[col_idx] = text.strip()
            context = headers.get(col_idx) or table_label
            yield (
                line,
                Unit(
                    text=text,
                    path=rel_path,
                    line=line,
                    column=f"{table_label}!{context}" if context != table_label else table_label,
                    context_key=context,
                    aggregate=False,
                ),
            )
            line += 1


class DocxAdapter:
    name = "docx"
    extensions = frozenset({".docx"})

    def supports(self, path: Path) -> bool:
        return path.suffix.lower() in self.extensions

    def iter_units(self, path: Path, *, rel_path: str) -> Iterator[Unit]:
        if not office_docx_available():
            _warn_missing_once()
            return
        from docx import Document
        from docx.opc.exceptions import PackageNotFoundError

        try:
            document = Document(str(path))
        except (OSError, PackageNotFoundError, ValueError, KeyError):
            return

        line = 1
        for next_line, unit in _iter_paragraph_units(
            document.paragraphs, rel_path=rel_path, context_key="body", line_start=line
        ):
            yield unit
            line = next_line + 1

        for table_idx, table in enumerate(document.tables, start=1):
            label = f"table{table_idx}"
            for next_line, unit in _iter_table_units(
                table, rel_path=rel_path, table_label=label, line_start=line
            ):
                yield unit
                line = next_line + 1

        for section_idx, section in enumerate(document.sections, start=1):
            for part_name, container in (
                ("header", section.header),
                ("footer", section.footer),
            ):
                context = f"section{section_idx}:{part_name}"
                for next_line, unit in _iter_paragraph_units(
                    container.paragraphs,
                    rel_path=rel_path,
                    context_key=context,
                    line_start=line,
                ):
                    yield unit
                    line = next_line + 1
                for table_idx, table in enumerate(container.tables, start=1):
                    label = f"{context}:table{table_idx}"
                    for next_line, unit in _iter_table_units(
                        table, rel_path=rel_path, table_label=label, line_start=line
                    ):
                        yield unit
                        line = next_line + 1
